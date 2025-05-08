#!/usr/bin/env python3
import os
import sys
import subprocess
from argparse import ArgumentParser

# Attempt to import pure-Python backend
PYHWP_AVAILABLE = False
try:
    from hwp5.filestructure import Hwp5File
    PYHWP_AVAILABLE = True
except ImportError:
    pass

# python-docx for post-conversion styling
PYDOCX_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    PYDOCX_AVAILABLE = True
except ImportError:
    pass


def convert_with_pyhwp(input_path, temp_output):
    if not PYHWP_AVAILABLE:
        print("ERROR: pyhwp not installed—cannot convert in python mode.", file=sys.stderr)
        sys.exit(1)
    hwp = Hwp5File(input_path)
    doc = Document()
    for rec in getattr(hwp, 'record_list', []):
        text = getattr(rec, 'text', '').strip()
        if text:
            doc.add_paragraph(text)
    doc.save(temp_output)


def convert_with_uno(input_path, temp_output):
    outdir = os.path.dirname(os.path.abspath(temp_output)) or '.'
    cmd = ['soffice', '--headless', '--convert-to', 'docx', '--outdir', outdir, input_path]
    try:
        subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except FileNotFoundError:
        print("ERROR: 'soffice' not found—install LibreOffice.", file=sys.stderr)
        sys.exit(1)
    generated = os.path.join(outdir, os.path.splitext(os.path.basename(input_path))[0] + '.docx')
    if os.path.abspath(generated) != os.path.abspath(temp_output):
        os.replace(generated, temp_output)


def adjust_font(input_docx, output_docx, font_name, size_delta, explicit_size):
    if not PYDOCX_AVAILABLE:
        print("ERROR: python-docx not installed—cannot adjust font.", file=sys.stderr)
        sys.exit(1)
    doc = Document(input_docx)

    # Adjust default Normal style
    try:
        style = doc.styles['Normal']
        if font_name:
            style.font.name = font_name
            style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if explicit_size is not None:
            style.font.size = Pt(explicit_size)
        elif style.font.size:
            style.font.size = Pt(max(style.font.size.pt - size_delta, 1))
    except Exception:
        pass

    # Helper to set run properties
    def process_run(run):
        if font_name:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if explicit_size is not None:
            run.font.size = Pt(explicit_size)
        elif size_delta and run.font.size:
            new_size = max(run.font.size.pt - size_delta, 1)
            run.font.size = Pt(new_size)

    # Apply to all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            process_run(run)

    # Apply to tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        process_run(run)

    doc.save(output_docx)


def main():
    parser = ArgumentParser(description="HWP → DOCX Converter with Font and Size Control")
    parser.add_argument('--mode', choices=['uno','python'], default='uno',
                        help="Conversion backend: 'uno' (LibreOffice) or 'python' (pyhwp text-only)")
    parser.add_argument('--font-name', help='Font name for all text runs (e.g. "Batang")')
    parser.add_argument('--size-delta', type=int, default=0,
                        help='Decrease existing font sizes by this many points')
    parser.add_argument('--font-size', type=int,
                        help='Explicitly set all font sizes to this point value')
    parser.add_argument('input', help='Source .hwp/.hwpx file')
    parser.add_argument('output', help='Target .docx file')
    args = parser.parse_args()

    if not os.path.isfile(args.input):
        print(f"ERROR: Input file not found: {args.input}", file=sys.stderr)
        sys.exit(1)

    do_adjust = args.font_name or args.size_delta or args.font_size is not None
    temp_output = args.output + '.tmp.docx' if do_adjust else args.output

    if args.mode == 'python':
        convert_with_pyhwp(args.input, temp_output)
    else:
        convert_with_uno(args.input, temp_output)

    if do_adjust:
        adjust_font(temp_output, args.output, args.font_name, args.size_delta, args.font_size)
        os.remove(temp_output)

    print(f"✅ Converted '{args.input}' → '{args.output}'")

if __name__ == '__main__':
    main()
